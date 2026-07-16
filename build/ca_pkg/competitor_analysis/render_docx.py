"""Word-report renderer (python-docx), IBM Plex Sans to match the scoring deck."""

from __future__ import annotations

from .models import CompetitorAnalysis

_FONT = "IBM Plex Sans"
_BLUE = (15, 98, 254)     # IBM blue
_GRAY = (82, 82, 82)


def _style(doc):
    from docx.shared import Pt

    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(10.5)


def _heading(doc, text: str, level: int):
    from docx.shared import RGBColor

    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = _FONT
        run.font.color.rgb = RGBColor(*_BLUE)
    return h


def render_docx(analysis: CompetitorAnalysis, path: str) -> str:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor

    doc = DocxDocument()
    _style(doc)

    _heading(doc, f"Competitor Bid-Strategy Analysis: {analysis.competitor}", 0)
    meta = doc.add_paragraph()
    run = meta.add_run(
        f"Focal vendor: {analysis.focal}   |   Generated: {analysis.generated_at}   |   "
        f"Procurements analyzed: {len(analysis.procurement_digests)}   |   "
        f"Documents analyzed: {analysis.docs_analyzed}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(*_GRAY)
    doc.add_paragraph(
        "Source: FOIA documents (proposals, pricing workbooks, evaluation scoresheets) "
        "from the SLED competitive-intelligence corpus."
    ).runs[0].font.size = Pt(9)

    _heading(doc, "Executive Summary", 1)
    doc.add_paragraph(analysis.executive_summary or "No summary produced.")

    for dim in analysis.dimensions:
        _heading(doc, dim.title, 1)
        for para in (dim.analysis or "").split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        if dim.evidence:
            _heading(doc, "Evidence", 2)
            for item in dim.evidence:
                p = doc.add_paragraph(style="List Bullet")
                if item.procurement:
                    p.add_run(f"{item.procurement}: ").bold = True
                p.add_run(item.detail)
        if dim.ibm_implications:
            _heading(doc, f"Implications for {analysis.focal}", 2)
            doc.add_paragraph(dim.ibm_implications)

    _heading(doc, "Appendix: Procurements Analyzed", 1)
    for digest in analysis.procurement_digests:
        p = doc.add_paragraph(style="List Bullet")
        label = digest.procurement
        extras = ", ".join(x for x in (digest.client, digest.year, digest.outcome) if x)
        p.add_run(label).bold = True
        if extras:
            p.add_run(f" ({extras})")
        if digest.source_docs:
            p.add_run(f" — {len(digest.source_docs)} docs")

    if analysis.warnings:
        _heading(doc, "Notes & Limitations", 1)
        for w in analysis.warnings:
            doc.add_paragraph(w, style="List Bullet")

    doc.save(path)
    return path
